import os
import shutil
from os import name
from bs4 import BeautifulSoup
from requests import get
from csv import writer
import csv
import json
from selenium import webdriver
import time
import docx
import PyPDF2
from docx2pdf import convert
import openpyxl


def stock_value(url,name):#function defination
    r = get(url)
    soup = BeautifulSoup(r.text, 'lxml')


    # get all tables
    tables = soup.find_all('table')

    # loop over each table
    for num,table in enumerate(tables):

        # create filename
        filename = name+'.csv' #naming the file

        # open file for writing
        with open(filename, 'w') as f:

            # store rows here
            data = []

            # create csv writer object
            csv_writer = writer(f)

            # go through each row
            rows = table.find_all('tr')
            for row in rows:

                # write headers if any
                headers = row.find_all('th')
                if headers:
                    try:
                        csv_writer.writerow([header.text.strip() for header in headers])
                    except:
                        pass
                # write column items
                columns = row.find_all('td')
                try:
                    csv_writer.writerow([column.text.strip() for column in columns])
                except:
                    pass
    return filename

def del_prev_stock_info():
    y=input("\nFiles related to previous execution will be deleted !\nAre you sure to continue execution in path : "+os.getcwd()+" [ Y / N ] : ")
    y=y.upper()
    if y=='Y':
        pass
    else:
        print("Please move and re-run this tool in a new folder ...")
        exit(0)
    z=os.listdir()
    for i in z:
        if(i=='test.py'):
            pass
        elif(i=='msedgedriver.exe'):
            pass
        elif(i=='chromedriver.exe'):
            pass
        elif(os.path.isdir(i)==True):
            pass
        else:
            os.remove(i)

def dir_ispresent(x):
    if os.path.isdir(x):
        shutil.rmtree(x)
    else:
        pass
    os.mkdir(x)
    
i = 1
while i == 1:
    del_prev_stock_info()
    print("\n*** KINDLY USE THIS TOOL IN A SEPARATE FOLDER ***\n\n\tMENU")
    print("1. Most active stocks")#menu
    print("2. Gainer stocks")
    print("3. Loser stocks")
    ch = int(input("\nEnter your choice : ")) #prompting user to choose
    if ch ==1:
        url = 'https://finance.yahoo.com/most-active'
        file_name ='active_stocks' 
        d = stock_value(url,file_name)

    if ch ==2:
        url2 = 'https://finance.yahoo.com/gainers'
        file_name = 'stock_gainers'
        d =stock_value(url2,file_name)
    if ch ==3:
        url3 = 'https://finance.yahoo.com/losers'
        file_name = 'stock_losers'
        d = stock_value(url3,file_name) #function call
    i +=1
    
    
csvRows = [] #list to store csv values
csvFileObj = open(d) #d contains the filename returned by the function value
readerObj = csv.reader(csvFileObj) #creating a reader object
for row in readerObj: #traversing through each row
    if readerObj.line_num == 1:
        continue
    csvRows.append(row) #appending row values
csvFileObj.close()
# print(csvRows)

with open(file_name+'.json','w') as jsonFile: #creating a jsonfile by name as that of file_name.
    json.dump(csvRows,jsonFile,indent=2) #dump function is used to write into the json file
    
def get_updates(url,option):#function defination
    browser = webdriver.Chrome()
    browser.get(url) 
    dropdown = browser.find_element_by_css_selector('#chart-toolbar > div:nth-child(1) > div.M\(0\).D\(ib\).Bdrs\(3px\).C\(\$tertiaryColor\).Fz\(s\).Fw\(500\).O\(n\)\:f\!.Mstart\(15px\).Mstart\(10px\)--tab1024.Mstart\(10px\)--tab768.chartTypeBtn.Pos\(r\) > div.Pos\(r\).D\(ib\).Cur\(p\) > svg.Mstart\(2px\).Va\(m\)\!.Mt\(-2px\).chartTypeBtn\:h_Fill\(\$linkColor\)\!.chartTypeBtn\:h_Stk\(\$linkColor\)\!.Cur\(p\)')
    dropdown.click() #it clicks the drop down menu
    ele = browser.find_element_by_css_selector(option) #option contains the value we chose for viewing the data 
    ele.click()
    time.sleep(5) #making the webdriver wait
    screenshot = browser.save_screenshot(stock+'.png') #this is used to take photos or screenshots of the graph
    elem = browser.find_element_by_css_selector('#chart-toolbar > div:nth-child(1) > div.D\(ib\).Whs\(nw\).H\(17px\) > ul > li:nth-child(2) > button > span') #selecting the 5 day change values
    elem.click()
    time.sleep(5)
    screenshot = browser.save_screenshot(stock+'5d.png') 
    elem2 = browser.find_element_by_css_selector('#chart-toolbar > div:nth-child(1) > div.D\(ib\).Whs\(nw\).H\(17px\) > ul > li:nth-child(3) > button > span') #selecting 1M changes values
    elem2.click()
    time.sleep(5)
    screenshot = browser.save_screenshot(stock+'1M.png')
    elem3 = browser.find_element_by_css_selector('#chart-toolbar > div:nth-child(1) > div.D\(ib\).Whs\(nw\).H\(17px\) > ul > li:nth-child(4) > button > span')#selecting 3M changed values
    elem3.click()
    time.sleep(5)
    screenshot = browser.save_screenshot(stock+'3M.png')
    elem4 = browser.find_element_by_css_selector('#chart-toolbar > div:nth-child(1) > div.D\(ib\).Whs\(nw\).H\(17px\) > ul > li:nth-child(5) > button > span')#selecting 6 month chnaged values
    elem4.click()
    time.sleep(5)
    screenshot = browser.save_screenshot(stock+'6M.png')
    elem5 = browser.find_element_by_css_selector('#chart-toolbar > div:nth-child(1) > div.D\(ib\).Whs\(nw\).H\(17px\) > ul > li:nth-child(7) > button > span')#selecting 1 year changed values
    elem5.click()
    time.sleep(5)
    screenshot = browser.save_screenshot(stock+'1Y.png') #saving the image as choosen stockname1y.png
    browser.quit() #closing the webdriver
    doc = docx.Document()
    doc.add_heading(f'{stock.upper()} details', 0) #adding headings to the document
    doc.add_picture(stock+'.png',width=docx.shared.Cm(17.71), height=docx.shared.Cm(12.7)) #inserting the captured images to word
    doc.add_page_break() #in order to make the image come in next line
    doc.add_heading(f'{stock.upper()} 5d change', 0) #heading 
    doc.add_picture(stock+'5d.png',width=docx.shared.Cm(17.71), height=docx.shared.Cm(12.7))
    doc.add_page_break()
    doc.add_heading(f'{stock.upper()} 1M change', 0)
    doc.add_picture(stock+'1M.png',width=docx.shared.Cm(17.71), height=docx.shared.Cm(12.7))
    doc.add_page_break()
    doc.add_heading(f'{stock.upper()} 3M change', 0)
    doc.add_picture(stock+'3M.png',width=docx.shared.Cm(17.71), height=docx.shared.Cm(12.7))
    doc.add_page_break()
    doc.add_heading(f'{stock.upper()} 6M change', 0)
    doc.add_picture(stock+'6M.png',width=docx.shared.Cm(17.71), height=docx.shared.Cm(12.7))
    doc.add_page_break()
    doc.add_heading(f'{stock.upper()} 1Y change', 0)
    doc.add_picture(stock+'1Y.png',width=docx.shared.Cm(17.71), height=docx.shared.Cm(12.7))
    doc.save(stock+'.docx')# saving the docx file
    return(stock+'.docx')# function returns selected stockfile.docx
    
while True:
    stock = input("\nEnter the stock name you want to see updates on : ") #prompting user to enter a stock value
    stock.upper()
    dir_ispresent(stock.upper())
    url4 = 'https://finance.yahoo.com/chart/'+stock+'#eyJpbnRlcnZhbCI6ImRheSIsInBlcmlvZGljaXR5IjoxLCJ0aW1lVW5pdCI6bnVsbCwiY2FuZGxlV2lkdGgiOjgsImZsaXBwZWQiOmZhbHNlLCJ2b2x1bWVVbmRlcmxheSI6dHJ1ZSwiYWRqIjp0cnVlLCJjcm9zc2hhaXIiOnRydWUsImNoYXJ0VHlwZSI6ImxpbmUiLCJleHRlbmRlZCI6ZmFsc2UsIm1hcmtldFNlc3Npb25zIjp7fSwiYWdncmVnYXRpb25UeXBlIjoib2hsYyIsImNoYXJ0U2NhbGUiOiJsaW5lYXIiLCJwYW5lbHMiOnsiY2hhcnQiOnsicGVyY2VudCI6MSwiZGlzcGxheSI6Ik5WREEiLCJjaGFydE5hbWUiOiJjaGFydCIsImluZGV4IjowLCJ5QXhpcyI6eyJuYW1lIjoiY2hhcnQiLCJwb3NpdGlvbiI6bnVsbH0sInlheGlzTEhTIjpbXSwieWF4aXNSSFMiOlsiY2hhcnQiLCLigIx2b2wgdW5kcuKAjCJdfX0sInNldFNwYW4iOnt9LCJsaW5lV2lkdGgiOjIsInN0cmlwZWRCYWNrZ3JvdW5kIjp0cnVlLCJldmVudHMiOnRydWUsImNvbG9yIjoiIzAwODFmMiIsInN0cmlwZWRCYWNrZ3JvdWQiOnRydWUsImV2ZW50TWFwIjp7ImNvcnBvcmF0ZSI6eyJkaXZzIjp0cnVlLCJzcGxpdHMiOnRydWV9LCJzaWdEZXYiOnt9fSwic3ltYm9scyI6W3sic3ltYm9sIjoiTlZEQSIsInN5bWJvbE9iamVjdCI6eyJzeW1ib2wiOiJOVkRBIiwicXVvdGVUeXBlIjoiRVFVSVRZIiwiZXhjaGFuZ2VUaW1lWm9uZSI6IkFtZXJpY2EvTmV3X1lvcmsifSwicGVyaW9kaWNpdHkiOjEsImludGVydmFsIjoiZGF5IiwidGltZVVuaXQiOm51bGwsInNldFNwYW4iOnt9fV0sInN0dWRpZXMiOnsi4oCMdm9sIHVuZHLigIwiOnsidHlwZSI6InZvbCB1bmRyIiwiaW5wdXRzIjp7ImlkIjoi4oCMdm9sIHVuZHLigIwiLCJkaXNwbGF5Ijoi4oCMdm9sIHVuZHLigIwifSwib3V0cHV0cyI6eyJVcCBWb2x1bWUiOiIjMDBiMDYxIiwiRG93biBWb2x1bWUiOiIjZmYzMzNhIn0sInBhbmVsIjoiY2hhcnQiLCJwYXJhbWV0ZXJzIjp7IndpZHRoRmFjdG9yIjowLjQ1LCJjaGFydE5hbWUiOiJjaGFydCIsInBhbmVsTmFtZSI6ImNoYXJ0In19fX0-'
    print("\nSelect an option")#asking user to enter the type in which they want to see the data
    print("1. Line")
    print("2. Area")
    print("3. Candle")
    print("4. Hollow candle")
    print("5. Bar")
    print("6. colored bar\n")
    choice1 = int(input("Enter your choice : "))
    if choice1 ==1:
        option = '#dropdown-menu > ul > li.Cur\(p\).C\(\$linkColor\)\:h.Bgc\(\$hoverBgColor\)\:h.chartBtn.Bgc\(\$hoverBgColor\) > button'
        docs = get_updates(url4,option) #option contains the css selector of the particular type just like if u select "line" then option contains the css selector of the "line"
    elif choice1 ==2:
        option = '#dropdown-menu > ul > li:nth-child(2)'
        docs = get_updates(url4,option)
    elif choice1 ==3:
        option = '#dropdown-menu > ul > li:nth-child(3)'
        docs = get_updates(url4,option)
    elif choice1 ==4:
        option = '#dropdown-menu > ul > li:nth-child(4)'
        docs = get_updates(url4,option)
    elif choice1 ==5:
        option = '#dropdown-menu > ul > li:nth-child(5)'
        docs = get_updates(url4,option)
    elif choice1 ==6:
        option = '#dropdown-menu > ul > li:nth-child(6)'
        docs = get_updates(url4,option)
    else:print(f"{ch} is an invalid option ")
    
    convert(docs) #converting docs to pdf using docx2pdf module
    try:
        convert('C:\\windows\\system32\\stock\\'+docs,'C:\\windows\\system32\\stock\\'+stock+'.pdf')
        convert("C:\\windows\\system32\\stock\\")
    except:
        pass
    csv_data = []
    with open(file_name+'.csv') as file_obj:
        reader = csv.reader(file_obj)
        for row in reader:
            csv_data.append(row)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    for row in csv_data:
        sheet.append(row)

    workbook.save(file_name+'.xlsx')
    choice = input("do you want to continue : ")
    if choice == 'y':
        for i in os.listdir():
            d=["test.py","msedgedriver.exe","chromedriver.exe","stock_losers.csv","active_stocks.csv","stock_gainers.csv","stock_losers.json","active_stocks.json","stock_gainers.json"]
            if i in d :
                pass
            else:
                shutil.move(i,'./'+stock)
        pass
    else:
        break

    
print("Thank you ..See you again !!")
